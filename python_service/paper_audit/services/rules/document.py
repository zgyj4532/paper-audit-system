from __future__ import annotations

import re
from typing import Any, Dict, Iterable, List

from .common import RuleIssue, get_section_format, get_section_text, get_sections
from .common import is_code_like_section

try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover - optional dependency guard
    DocxDocument = None


_BODY_FONT_NAMES = ("宋体", "仿宋", "simsun", "fangsong")


def _explicit_section_position(
    section: Dict[str, Any], section_index: int
) -> Dict[str, Any]:
    position = section.get("position")
    result: Dict[str, Any] = {}
    if isinstance(position, dict):
        for key in (
            "section_id",
            "paragraph_index",
            "page_start",
            "page_end",
            "xml_path",
        ):
            value = position.get(key)
            if value is not None:
                result[key] = value
    if "section_id" not in result:
        result["section_id"] = section.get("id", section_index)
    return result


def _explicit_table_meta(section: Dict[str, Any]) -> Dict[str, Any]:
    table_meta = section.get("table_meta")
    return table_meta if isinstance(table_meta, dict) else {}


def _table_position_from_section(
    section: Dict[str, Any], section_index: int
) -> Dict[str, Any]:
    position = _explicit_section_position(section, section_index)
    table_meta = _explicit_table_meta(section)
    for key in ("page_start", "page_end"):
        value = table_meta.get(key)
        if value is not None and key not in position:
            position[key] = value
    return position


def _load_document(source_file: str | None) -> Any | None:
    if not source_file or DocxDocument is None:
        return None
    try:
        return DocxDocument(source_file)
    except Exception:
        return None


def _paragraph_text(paragraph: Any) -> str:
    return str(getattr(paragraph, "text", "") or "").strip()


def _paragraph_font_name(paragraph: Any) -> str:
    for run in getattr(paragraph, "runs", []) or []:
        font_name = getattr(getattr(run, "font", None), "name", None)
        if font_name:
            return str(font_name)
    style = getattr(paragraph, "style", None)
    if style is not None:
        font_name = getattr(getattr(style, "font", None), "name", None)
        if font_name:
            return str(font_name)
    return ""


def _paragraph_font_size_pt(paragraph: Any) -> float | None:
    for run in getattr(paragraph, "runs", []) or []:
        font_size = getattr(getattr(run, "font", None), "size", None)
        if font_size is not None:
            try:
                return float(font_size.pt)
            except Exception:
                continue
    style = getattr(paragraph, "style", None)
    if style is not None:
        font_size = getattr(getattr(style, "font", None), "size", None)
        if font_size is not None:
            try:
                return float(font_size.pt)
            except Exception:
                return None
    return None


def _paragraph_line_spacing(paragraph: Any) -> float | None:
    paragraph_format = getattr(paragraph, "paragraph_format", None)
    if paragraph_format is not None:
        line_spacing = getattr(paragraph_format, "line_spacing", None)
        if isinstance(line_spacing, (int, float)):
            return float(line_spacing)
        if line_spacing is not None:
            try:
                return float(line_spacing.pt) / 12.0
            except Exception:
                return None
    return None


def _paragraph_first_line_indent_inch(paragraph: Any) -> float | None:
    paragraph_format = getattr(paragraph, "paragraph_format", None)
    if paragraph_format is None:
        return None
    first_line_indent = getattr(paragraph_format, "first_line_indent", None)
    if first_line_indent is None:
        return None
    try:
        return float(first_line_indent.inches)
    except Exception:
        return None


def _paragraph_left_indent_inch(paragraph: Any) -> float | None:
    paragraph_format = getattr(paragraph, "paragraph_format", None)
    if paragraph_format is None:
        return None
    left_indent = getattr(paragraph_format, "left_indent", None)
    if left_indent is None:
        return None
    try:
        return float(left_indent.inches)
    except Exception:
        return None


def _style_name(paragraph: Any) -> str:
    style = getattr(paragraph, "style", None)
    return str(getattr(style, "name", "") or getattr(style, "style_id", "") or "")


def _heading_level_from_style(style_name: str) -> int | None:
    normalized = re.sub(r"\s+", "", style_name).lower()
    for level in (1, 2, 3):
        if f"heading{level}" in normalized or f"标题{level}" in normalized:
            return level
    return None


def _heading_level_from_text(text: str) -> int | None:
    if re.match(r"^\s*[1-9]\d*[.．、]\s+", text):
        return 1
    if re.match(r"^\s*[1-9]\d*\.[1-9]\d*[.．、]\s+", text):
        return 2
    if re.match(r"^\s*[1-9]\d*\.[1-9]\d*\.[1-9]\d*[.．、]\s+", text):
        return 3
    if re.match(r"^\s*第[一二三四五六七八九十百千0-9]+[章节篇部分]", text):
        return 1
    if re.match(r"^\s*[一二三四五六七八九十]+[、.．]\s*", text):
        return 1
    if re.match(r"^\s*[（(][一二三四五六七八九十]+[)）]\s*", text):
        return 2
    return None


def _heading_level_for_section(section: Dict[str, Any]) -> int | None:
    format_info = get_section_format(section)
    style_name = str(format_info.get("paragraph_style") or "")
    text = get_section_text(section)
    return _heading_level_from_style(style_name) or _heading_level_from_text(text)


def _detect_caption(text: str, kind: str) -> bool:
    if not text:
        return False
    if kind == "table":
        return bool(re.search(r"^(?:表\s*\d+|Table\s*\d+)", text, re.IGNORECASE))
    return bool(
        re.search(r"^(?:图\s*\d+|Fig\.?\s*\d+|Figure\s*\d+)", text, re.IGNORECASE)
    )


def _add_issue(
    issues: List[Dict[str, Any]],
    *,
    issue_type: str,
    severity: int,
    message: str,
    suggestion: str,
    rule_id: str,
    position: Dict[str, Any] | None = None,
    original: str | None = None,
) -> None:
    issues.append(
        RuleIssue(
            issue_type=issue_type,
            severity=severity,
            message=message,
            suggestion=suggestion,
            position=position,
            original=original,
            rule_id=rule_id,
        ).as_dict()
    )


def _parsed_sections_rules(
    sections: List[Dict[str, Any]], issues: List[Dict[str, Any]]
) -> None:
    previous_heading_level: int | None = None
    previous_text = ""

    for index, section in enumerate(sections, start=1):
        if is_code_like_section(section):
            continue
        text = get_section_text(section)
        format_info = get_section_format(section)
        heading_level = _heading_level_for_section(section)
        section_position = _explicit_section_position(section, index)
        table_meta = _explicit_table_meta(section)

        if heading_level is not None:
            font_name = str(format_info.get("font") or "")
            size_raw = format_info.get("size")
            try:
                font_size = (
                    float(str(size_raw).replace("pt", ""))
                    if size_raw not in (None, "unknown", "")
                    else None
                )
            except Exception:
                font_size = None

            if heading_level == 1:
                if font_name and not any(
                    marker.lower() in font_name.lower() for marker in ("黑体", "simhei")
                ):
                    _add_issue(
                        issues,
                        issue_type="title_font",
                        severity=3,
                        message="一级标题字体不符合要求",
                        suggestion="一级标题应使用黑体",
                        rule_id="TITLE-001",
                        position=section_position,
                        original=font_name or None,
                    )
                if font_size is not None and not (17.0 <= font_size <= 19.0):
                    _add_issue(
                        issues,
                        issue_type="title_size",
                        severity=3,
                        message="一级标题字号不符合要求",
                        suggestion="一级标题应为 18pt±1pt",
                        rule_id="TITLE-002",
                        position=section_position,
                        original=str(font_size),
                    )
            elif heading_level == 2:
                if font_name and not any(
                    marker.lower() in font_name.lower() for marker in ("黑体", "simhei")
                ):
                    _add_issue(
                        issues,
                        issue_type="title_font",
                        severity=3,
                        message="二级标题字体不符合要求",
                        suggestion="二级标题应使用黑体",
                        rule_id="TITLE-003",
                        position=section_position,
                        original=font_name or None,
                    )
                if font_size is not None and not (15.0 <= font_size <= 17.0):
                    _add_issue(
                        issues,
                        issue_type="title_size",
                        severity=3,
                        message="二级标题字号不符合要求",
                        suggestion="二级标题应为 16pt±1pt",
                        rule_id="TITLE-004",
                        position=section_position,
                        original=str(font_size),
                    )
            elif heading_level == 3:
                if font_name and not any(
                    marker.lower() in font_name.lower() for marker in ("黑体", "simhei")
                ):
                    _add_issue(
                        issues,
                        issue_type="title_font",
                        severity=3,
                        message="三级标题字体不符合要求",
                        suggestion="三级标题应使用黑体",
                        rule_id="TITLE-005",
                        position=section_position,
                        original=font_name or None,
                    )
                if font_size is not None and not (13.0 <= font_size <= 15.0):
                    _add_issue(
                        issues,
                        issue_type="title_size",
                        severity=3,
                        message="三级标题字号不符合要求",
                        suggestion="三级标题应为 14pt±1pt",
                        rule_id="TITLE-006",
                        position=section_position,
                        original=str(font_size),
                    )

            if previous_heading_level is not None:
                if heading_level > previous_heading_level + 1:
                    _add_issue(
                        issues,
                        issue_type="title_hierarchy",
                        severity=3,
                        message="标题层级跳跃",
                        suggestion="按 H1→H2→H3 逐级展开",
                        rule_id="TITLE-007",
                        position=section_position,
                        original=f"{previous_heading_level}->{heading_level}",
                    )
                if heading_level + 1 < previous_heading_level:
                    _add_issue(
                        issues,
                        issue_type="title_hierarchy",
                        severity=2,
                        message="标题层级出现逆序",
                        suggestion="检查标题编号与层级是否匹配",
                        rule_id="TITLE-008",
                        position=section_position,
                        original=f"{previous_heading_level}->{heading_level}",
                    )
            previous_heading_level = heading_level
            previous_text = text
            continue

        if bool(section.get("is_table")):
            table_position = _table_position_from_section(section, index)
            table_is_cross_page = bool(table_meta.get("is_cross_page")) or (
                isinstance(table_meta.get("page_start"), int)
                and isinstance(table_meta.get("page_end"), int)
                and int(table_meta["page_end"]) > int(table_meta["page_start"])
            )

            if not _detect_caption(previous_text, "table"):
                _add_issue(
                    issues,
                    issue_type="table_caption",
                    severity=2,
                    message="表格缺少题注",
                    suggestion="在表格前添加题注",
                    rule_id="TABLE-CAPTION-001",
                    position=table_position,
                    original=text[:80] or None,
                )
            if table_is_cross_page or (
                len(section.get("table_rows") or []) > 15
                and "续表" not in previous_text
            ):
                _add_issue(
                    issues,
                    issue_type="table_continuation",
                    severity=2,
                    message="表格可能跨页但未设置续表标志",
                    suggestion="跨页表格应标注“续表”",
                    rule_id="TABLE-CONT-001",
                    position=table_position,
                    original=text[:80] or None,
                )
            previous_text = text
            continue

        if section.get("images") and not _detect_caption(previous_text, "image"):
            _add_issue(
                issues,
                issue_type="image_caption",
                severity=2,
                message="图片缺少题注",
                suggestion="在图片前后补充题注",
                rule_id="IMAGE-CAPTION-001",
                position=section_position,
                original=text[:80] or None,
            )

        if not text.strip():
            _add_issue(
                issues,
                issue_type="blank_lines",
                severity=2,
                message="存在多余空行",
                suggestion="删除连续空段落",
                rule_id="SPACE-001",
                position=section_position,
            )
        previous_text = text


def _section_page_number_rule(
    section: Any, section_index: int, issues: List[Dict[str, Any]]
) -> None:
    footer = getattr(section, "footer", None)
    if footer is None:
        _add_issue(
            issues,
            issue_type="page_number",
            severity=2,
            message="页码未居中",
            suggestion="页码应居中放置",
            rule_id="PAGE-001",
            position={"section": section_index},
        )
        return

    footer_xml = str(getattr(getattr(footer, "_element", None), "xml", "") or "")
    if "PAGE" not in footer_xml.upper():
        _add_issue(
            issues,
            issue_type="page_number",
            severity=2,
            message="页码未居中",
            suggestion="页码应居中放置",
            rule_id="PAGE-001",
            position={"section": section_index},
        )
        return

    if (
        'w:jc w:val="center"' not in footer_xml
        and "w:jc w:val='center'" not in footer_xml
    ):
        _add_issue(
            issues,
            issue_type="page_number",
            severity=2,
            message="页码未居中",
            suggestion="页码应居中放置",
            rule_id="PAGE-001",
            position={"section": section_index},
        )


def _section_margin_rule(
    section: Any, section_index: int, issues: List[Dict[str, Any]]
) -> None:
    margins = {
        "top": getattr(getattr(section, "top_margin", None), "cm", None),
        "bottom": getattr(getattr(section, "bottom_margin", None), "cm", None),
        "left": getattr(getattr(section, "left_margin", None), "cm", None),
        "right": getattr(getattr(section, "right_margin", None), "cm", None),
    }
    for margin_name, margin_value in margins.items():
        if margin_value is None:
            continue
        try:
            numeric_value = float(margin_value)
        except Exception:
            continue
        if not (2.2 <= numeric_value <= 2.8):
            _add_issue(
                issues,
                issue_type="page_margin",
                severity=3,
                message="页边距不符合要求",
                suggestion="页边距应为 2.5cm±0.3cm",
                rule_id="MARGIN-001",
                position={"section": section_index},
                original=f"{margin_name}={numeric_value:.2f}cm",
            )


def _iter_docx_blocks(document: Any) -> Iterable[tuple[str, Any]]:
    paragraph_index = 0
    table_index = 0
    for child in document.element.body.iterchildren():
        tag = getattr(child, "tag", "")
        if tag.endswith("}p"):
            if paragraph_index < len(document.paragraphs):
                yield "paragraph", document.paragraphs[paragraph_index]
            paragraph_index += 1
        elif tag.endswith("}tbl"):
            if table_index < len(document.tables):
                yield "table", document.tables[table_index]
            table_index += 1


def _docx_paragraph_rules(document: Any, issues: List[Dict[str, Any]]) -> None:
    previous_heading_level: int | None = None
    blank_count = 0

    for index, (kind, block) in enumerate(_iter_docx_blocks(document), start=1):
        if kind != "paragraph":
            continue

        text = _paragraph_text(block)
        if not text:
            blank_count += 1
            if blank_count > 1:
                _add_issue(
                    issues,
                    issue_type="blank_lines",
                    severity=2,
                    message="存在多余空行",
                    suggestion="删除连续空段落",
                    rule_id="SPACE-001",
                    position={"paragraph": index},
                )
            continue
        blank_count = 0

        heading_level = _heading_level_from_style(
            _style_name(block)
        ) or _heading_level_from_text(text)
        if heading_level is not None:
            line_spacing = _paragraph_line_spacing(block)
            if line_spacing is not None and not (1.45 <= line_spacing <= 1.55):
                _add_issue(
                    issues,
                    issue_type="line_spacing",
                    severity=3,
                    message="行距不符合要求",
                    suggestion="行距应为 1.5 倍±0.05",
                    rule_id="LINE-001",
                    position={"paragraph": index},
                    original=f"{line_spacing:.2f}",
                )

            next_paragraph = None
            for next_kind, next_block in list(_iter_docx_blocks(document))[index:]:
                if next_kind == "paragraph" and _paragraph_text(next_block):
                    next_paragraph = next_block
                    break
            if next_paragraph is not None:
                spacing_value = None
                paragraph_format = getattr(block, "paragraph_format", None)
                if paragraph_format is not None:
                    space_after = getattr(paragraph_format, "space_after", None)
                    if space_after is not None:
                        try:
                            spacing_value = float(space_after.pt)
                        except Exception:
                            spacing_value = None
                if spacing_value is None:
                    next_format = getattr(next_paragraph, "paragraph_format", None)
                    if next_format is not None:
                        space_before = getattr(next_format, "space_before", None)
                        if space_before is not None:
                            try:
                                spacing_value = float(space_before.pt)
                            except Exception:
                                spacing_value = None
                if spacing_value is not None and not (3.6 <= spacing_value <= 8.4):
                    _add_issue(
                        issues,
                        issue_type="title_body_spacing",
                        severity=2,
                        message="标题与正文间距不符合要求",
                        suggestion="标题与正文间距应为 0.5 行（约 0.3~0.7）",
                        rule_id="SPACE-002",
                        position={"paragraph": index},
                        original=f"{spacing_value:.2f}pt",
                    )

            font_name = _paragraph_font_name(block)
            if (
                font_name
                and heading_level == 1
                and not any(
                    marker.lower() in font_name.lower() for marker in ("黑体", "simhei")
                )
            ):
                _add_issue(
                    issues,
                    issue_type="title_font",
                    severity=3,
                    message="一级标题字体不符合要求",
                    suggestion="一级标题应使用黑体",
                    rule_id="TITLE-001",
                    position={"paragraph": index},
                    original=font_name,
                )
            if (
                font_name
                and heading_level == 2
                and not any(
                    marker.lower() in font_name.lower() for marker in ("黑体", "simhei")
                )
            ):
                _add_issue(
                    issues,
                    issue_type="title_font",
                    severity=3,
                    message="二级标题字体不符合要求",
                    suggestion="二级标题应使用黑体",
                    rule_id="TITLE-003",
                    position={"paragraph": index},
                    original=font_name,
                )
            if (
                font_name
                and heading_level == 3
                and not any(
                    marker.lower() in font_name.lower() for marker in ("黑体", "simhei")
                )
            ):
                _add_issue(
                    issues,
                    issue_type="title_font",
                    severity=3,
                    message="三级标题字体不符合要求",
                    suggestion="三级标题应使用黑体",
                    rule_id="TITLE-005",
                    position={"paragraph": index},
                    original=font_name,
                )

            font_size = _paragraph_font_size_pt(block)
            if font_size is not None:
                if heading_level == 1 and not (17.0 <= font_size <= 19.0):
                    _add_issue(
                        issues,
                        issue_type="title_size",
                        severity=3,
                        message="一级标题字号不符合要求",
                        suggestion="一级标题应为 18pt±1pt",
                        rule_id="TITLE-002",
                        position={"paragraph": index},
                        original=f"{font_size:.2f}",
                    )
                if heading_level == 2 and not (15.0 <= font_size <= 17.0):
                    _add_issue(
                        issues,
                        issue_type="title_size",
                        severity=3,
                        message="二级标题字号不符合要求",
                        suggestion="二级标题应为 16pt±1pt",
                        rule_id="TITLE-004",
                        position={"paragraph": index},
                        original=f"{font_size:.2f}",
                    )
                if heading_level == 3 and not (13.0 <= font_size <= 15.0):
                    _add_issue(
                        issues,
                        issue_type="title_size",
                        severity=3,
                        message="三级标题字号不符合要求",
                        suggestion="三级标题应为 14pt±1pt",
                        rule_id="TITLE-006",
                        position={"paragraph": index},
                        original=f"{font_size:.2f}",
                    )

            if previous_heading_level is not None:
                if heading_level > previous_heading_level + 1:
                    _add_issue(
                        issues,
                        issue_type="title_hierarchy",
                        severity=3,
                        message="标题层级跳跃",
                        suggestion="按 H1→H2→H3 逐级展开",
                        rule_id="TITLE-007",
                        position={"paragraph": index},
                        original=f"{previous_heading_level}->{heading_level}",
                    )
                if heading_level + 1 < previous_heading_level:
                    _add_issue(
                        issues,
                        issue_type="title_hierarchy",
                        severity=2,
                        message="标题层级出现逆序",
                        suggestion="检查标题编号与层级是否匹配",
                        rule_id="TITLE-008",
                        position={"paragraph": index},
                        original=f"{previous_heading_level}->{heading_level}",
                    )
            previous_heading_level = heading_level
        else:
            font_name = _paragraph_font_name(block)
            font_size = _paragraph_font_size_pt(block)
            if font_name and not any(
                marker.lower() in font_name.lower() for marker in _BODY_FONT_NAMES
            ):
                _add_issue(
                    issues,
                    issue_type="body_font",
                    severity=3,
                    message="正文字体不符合要求",
                    suggestion="正文应使用宋体或仿宋",
                    rule_id="BODY-001",
                    position={"paragraph": index},
                    original=font_name,
                )
            if font_size is not None and not (11.0 <= font_size <= 13.0):
                _add_issue(
                    issues,
                    issue_type="body_size",
                    severity=3,
                    message="正文字号不符合要求",
                    suggestion="正文应为 12pt±1pt",
                    rule_id="BODY-002",
                    position={"paragraph": index},
                    original=f"{font_size:.2f}",
                )

            indent_value = _paragraph_first_line_indent_inch(block)
            if indent_value is not None and not (0.4 <= indent_value <= 0.6):
                _add_issue(
                    issues,
                    issue_type="body_indent",
                    severity=2,
                    message="段落首行缩进不符合要求",
                    suggestion="首行缩进应为 2 个字符（约 0.4~0.6 英寸）",
                    rule_id="BODY-003",
                    position={"paragraph": index},
                    original=f"{indent_value:.2f}",
                )

            left_indent = _paragraph_left_indent_inch(block)
            if left_indent is not None:
                list_like = bool(
                    re.match(
                        r"^(?:\d+\.|[一二三四五六七八九十]+[、.．]|[（(][一二三四五六七八九十]+[)）])",
                        text,
                    )
                    or "list" in _style_name(block).lower()
                )
                if list_like:
                    level_match = re.search(r"(\d+)$", _style_name(block))
                    level = int(level_match.group(1)) if level_match else 1
                    expected_indent = 0.5 * level
                    if abs(left_indent - expected_indent) > 0.2:
                        _add_issue(
                            issues,
                            issue_type="list_indent",
                            severity=2,
                            message="列表项缩进不按层级递增",
                            suggestion="每层列表缩进应按 0.5 英寸递增",
                            rule_id="LIST-001",
                            position={"paragraph": index},
                            original=f"{left_indent:.2f}",
                        )


def _docx_layout_rules(document: Any, issues: List[Dict[str, Any]]) -> None:
    for section_index, section in enumerate(document.sections, start=1):
        _section_margin_rule(section, section_index, issues)
        _section_page_number_rule(section, section_index, issues)


def check_document_rules(
    parsed_data: Dict[str, Any], source_file: str | None = None
) -> List[Dict[str, Any]]:
    issues: List[Dict[str, Any]] = []
    sections = get_sections(parsed_data)

    _parsed_sections_rules(sections, issues)

    for index, section in enumerate(sections, start=1):
        if is_code_like_section(section):
            continue
        if bool(section.get("has_math")):
            alignment = str(get_section_format(section).get("alignment") or "").lower()
            if alignment != "right":
                _add_issue(
                    issues,
                    issue_type="formula_alignment",
                    severity=2,
                    message="公式未右对齐",
                    suggestion="公式应右对齐",
                    rule_id="FORMULA-001",
                    position=_explicit_section_position(section, index),
                    original=alignment or None,
                )

    document = _load_document(source_file)
    if document is not None:
        _docx_paragraph_rules(document, issues)
        _docx_layout_rules(document, issues)

    return issues
