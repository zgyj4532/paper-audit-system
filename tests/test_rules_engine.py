import pathlib
import sys
from pathlib import Path

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parents[1]))

from docx import Document
from docx.shared import Cm, Inches, Pt

from python_service.paper_audit.services.rules.engine import (
    check_consistency_rules,
    check_document_rules,
    check_text_rules,
)


def test_gbt_format_rules_cover_common_thesis_defects():
    text = (
        "Abstract:This paper studies panoramic simulation.\n"
        "Keywords:VirtualReality;ImageStitchingFusion; Unity3D;PanoramaTechnology\n"
        "杨笛航. 全景视频图像融合与拼接算法研究D]. 浙江:浙江大学,2017.\n"
        "彭凤婷. 基于多全景相机拼接的虚拟现实和实景交互系统[[D]. 四川:电子科技大学,2017."
    )

    issues = check_text_rules(text, focus_areas=("format", "reference"))
    rule_ids = {issue.get("rule_id") for issue in issues}

    assert "FORMAT-004" in rule_ids
    assert "FORMAT-006" in rule_ids
    assert "FORMAT-005" in rule_ids
    assert "REF-002" in rule_ids
    assert "REF-003" in rule_ids


def test_gbt_rules_still_detect_reference_section_problems():
    text = (
        "参考文献\n"
        "杨笛航. 全景视频图像融合与拼接算法研究D]. 浙江:浙江大学,2017.\n"
        "彭凤婷. 基于多全景相机拼接的虚拟现实和实景交互系统[[D]. 四川:电子科技大学,2017."
    )

    issues = check_text_rules(text, focus_areas=("reference",))
    assert any(issue.get("issue_type") == "reference" for issue in issues)


def test_consistency_rules_detects_unexpanded_abbreviations_and_skips_reference_and_code():
    parsed_data = {
        "sections": [
            {
                "id": 1,
                "element_type": "Paragraph",
                "raw_text": "本文研究了CNN在图像分类中的应用。",
                "formatting": {"paragraph_style": "Normal"},
            },
            {
                "id": 2,
                "element_type": "Paragraph",
                "raw_text": "参考文献 [1] Smith J. CNN in vision. 2024.",
                "formatting": {"paragraph_style": "Normal"},
            },
            {
                "id": 3,
                "element_type": "Paragraph",
                "raw_text": (
                    "using System.Collections.Generic;\n"
                    "public class Demo {\n"
                    "    public void Run() {\n"
                    "        camera.DOFieldOfView(60, 1);\n"
                    "    }\n"
                    "}"
                ),
                "formatting": {"paragraph_style": "Normal"},
            },
        ]
    }

    issues = check_consistency_rules(parsed_data)
    rule_ids = [issue.get("rule_id") for issue in issues]

    assert rule_ids.count("CONSIST-003") == 1
    assert any(issue.get("original") == "CNN" for issue in issues)


def test_consistency_rules_skips_pure_english_articles():
    parsed_data = {
        "sections": [
            {
                "id": 1,
                "element_type": "Paragraph",
                "raw_text": (
                    "This paper studies convolutional neural networks and image classification.\n"
                    "No Chinese text appears in this article."
                ),
                "formatting": {"paragraph_style": "Normal"},
            }
        ]
    }

    issues = check_consistency_rules(parsed_data)
    assert all(issue.get("rule_id") != "CONSIST-003" for issue in issues)


def test_document_and_reference_rules_cover_layout_and_reference_gaps(tmp_path: Path):
    docx_path = tmp_path / "layout_rules.docx"
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

    heading1 = document.add_paragraph("1 绪论")
    heading1.style = document.styles["Heading 1"]
    heading1.runs[0].font.name = "Arial"
    heading1.runs[0].font.size = Pt(16)
    heading1.paragraph_format.line_spacing = 1.0
    heading1.paragraph_format.first_line_indent = Inches(0.1)

    heading3 = document.add_paragraph("1.1 背景")
    heading3.style = document.styles["Heading 3"]
    heading3.runs[0].font.name = "Arial"
    heading3.runs[0].font.size = Pt(13)
    heading3.paragraph_format.line_spacing = 1.0

    body = document.add_paragraph("正文内容")
    body.runs[0].font.name = "Arial"
    body.runs[0].font.size = Pt(10)
    body.paragraph_format.line_spacing = 1.0
    body.paragraph_format.first_line_indent = Inches(0.1)

    document.add_paragraph("")
    document.add_paragraph("")

    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"

    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.text = "1"
    footer_paragraph.alignment = 0

    document.save(docx_path)

    parsed_data = {
        "sections": [
            {
                "id": 1,
                "element_type": "Paragraph",
                "position": {
                    "section_id": 1,
                    "paragraph_index": 1,
                    "page_start": 1,
                    "page_end": 1,
                    "xml_path": "/w:body/w:p[1]",
                },
                "raw_text": "1 绪论",
                "formatting": {
                    "font": "Arial",
                    "size": "16pt",
                    "alignment": "left",
                    "paragraph_style": "Heading 1",
                },
            },
            {
                "id": 2,
                "element_type": "Paragraph",
                "position": {
                    "section_id": 2,
                    "paragraph_index": 2,
                    "page_start": 1,
                    "page_end": 1,
                    "xml_path": "/w:body/w:p[2]",
                },
                "raw_text": "1.1 背景",
                "formatting": {
                    "font": "Arial",
                    "size": "13pt",
                    "alignment": "left",
                    "paragraph_style": "Heading 3",
                },
            },
            {
                "id": 3,
                "element_type": "Paragraph",
                "raw_text": "正文内容",
                "formatting": {
                    "font": "Arial",
                    "size": "10pt",
                    "alignment": "left",
                    "paragraph_style": "Normal",
                },
            },
            {
                "id": 4,
                "element_type": "Table",
                "raw_text": "A\tB",
                "is_table": True,
                "table_rows": [["A", "B"]],
                "table_meta": {
                    "row_count": 1,
                    "column_count": 2,
                    "page_start": 2,
                    "page_end": 3,
                    "is_cross_page": True,
                },
                "has_math": True,
                "formatting": {"font": "宋体", "size": "12pt", "alignment": "left"},
            },
            {
                "id": 5,
                "element_type": "Paragraph",
                "raw_text": "文中引用[1]和[5]。",
                "formatting": {
                    "font": "宋体",
                    "size": "12pt",
                    "alignment": "left",
                    "paragraph_style": "Normal",
                },
            },
        ],
        "references": [
            {"text": "[1] 张三,李四,王五. 题名[J]，期刊, 2027, 35(6): 18-21."},
            {"text": "[4] 张三. 另一篇期刊论文[J]. 某期刊, 98, 35(6): 18-21."},
            {"text": "[2] 赵六. 学位论文题名[M]. 北京: 某大学, 98."},
            {"text": "[3] 王五. 另一篇硕士论文[M]. 北京: 某大学, 1899."},
        ],
    }

    issues = check_consistency_rules(parsed_data, source_file=str(docx_path))
    rule_ids = {issue.get("rule_id") for issue in issues}
    # issues_by_rule = {
    #     issue.get("rule_id"): issue for issue in issues if isinstance(issue, dict)
    # }
    document_issues = check_document_rules(parsed_data)
    document_issues_by_rule = {
        issue.get("rule_id"): issue
        for issue in document_issues
        if isinstance(issue, dict)
    }

    assert "LINE-001" in rule_ids
    assert "TITLE-007" in rule_ids
    assert "TITLE-001" in rule_ids
    assert "TITLE-002" in rule_ids
    assert "BODY-001" in rule_ids
    assert "BODY-002" in rule_ids
    assert "BODY-003" in rule_ids
    assert "TABLE-CAPTION-001" in rule_ids
    assert "FORMULA-001" in rule_ids
    assert "MARGIN-001" in rule_ids
    assert "PAGE-001" in rule_ids
    assert "SPACE-001" in rule_ids
    assert "REF-J-001" in rule_ids
    assert "REF-J-003" in rule_ids
    assert "REF-J-007" in rule_ids
    assert "REF-J-011" in rule_ids
    assert "REF-M-001" in rule_ids
    assert "REF-M-002" in rule_ids
    assert "REF-CONSIST-001" in rule_ids
    assert "REF-CONSIST-002" in rule_ids

    title_issue = document_issues_by_rule.get("TITLE-007")
    assert title_issue is not None
    assert title_issue.get("position", {}).get("paragraph_index") == 2
    assert title_issue.get("position", {}).get("page_start") == 1

    table_issue = document_issues_by_rule.get("TABLE-CONT-001")
    assert table_issue is not None
    assert table_issue.get("position", {}).get("page_start") == 2
    assert table_issue.get("position", {}).get("page_end") == 3
